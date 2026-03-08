import docx
import os

def extract_text_from_docx(file_path):
    if not os.path.exists(file_path):
        return f"File not found: {file_path}"
    
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading docx: {str(e)}"

if __name__ == "__main__":
    file_path = r"e:\projects\cricmatch_predict\CricPredict_Project_Plan.docx"
    content = extract_text_from_docx(file_path)
    with open(r"e:\projects\cricmatch_predict\project_plan.txt", "w", encoding="utf-8") as f:
        f.write(content)
    print("Successfully wrote to project_plan.txt")
